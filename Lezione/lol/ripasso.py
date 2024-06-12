from docx import Document
from docx.shared import Inches
import pyttsx3

pyttsx3.speak('Hello')


document = Document()

#  profile picture
document.add_picture(
    "bobdylan.jpg",
    width= Inches(2.0)
)

#  name phone number and email details
name = input("What is your name? ")
number_phone = input("What is your phone number? ")
email = input("What is your email? ")


document.add_paragraph(
    f'{name} | {number_phone} | {email}'
) 

# about me
document.add_heading("About me")
document.add_paragraph(
    input("Tell about yourself? ")
)

# work experience
document.add_heading("Work experience")
p = document.add_paragraph()

company = input('Enter Company ')
from_date = input('From Date ')
to_date = input('To Date ')

p.add_run(company + ' ').bold = True
p.add_run(from_date + '-' + to_date + '\n').italic = True


experiance_details = input(f'Experienc in {company} ')
p.add_run(experiance_details)

# more experiences
while True:
    more_experinces = input(
        'Do you have more experinces? Yes or No ')
    if more_experinces.lower() == 'yes':
        p = document.add_paragraph()

        company = input('Enter Company ')
        from_date = input('From Date ')
        to_date = input('To Date ')

        p.add_run(company + ' ').bold = True
        p.add_run(from_date + '-' + to_date + '\n').italic = True

        experiance_details = input(f'Experienc in {company} ')
        p.add_run(experiance_details)
    else:
        break

# personal skills
document.add_heading("Skills")
skill = input('Enter skill ')
p = document.add_paragraph(skill)
p.style = 'List Bullet'

while True:
    more_skills = input('You have another skills? Yes or No ')
    if more_skills.lower() == 'yes':
        skill = input('Enter skill ')
        p = document.add_paragraph(skill)
        p.style = 'List Bullet'
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = 'CV generated using Amigoscode and Intuiut QuickBooks'

document.save('cv.docx')